from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
from docx.shared import Pt
import pandas as pd
import numpy as np
import os
import shutil
from tqdm import tqdm
import glob
from zipfile import ZipFile 
import concurrent
pd.options.mode.chained_assignment = None 
import json
from docx.shared import RGBColor


from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tqdm import tqdm

def set_cell_text(cell, text, font_size=9, font_name='Arial', font_color=None, bold=False, align=None):
    """Set cell text with font formatting."""
    cell.text = (text or '').strip()
    for paragraph in cell.paragraphs:
        pPr = paragraph._p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:before'), '0')
        pPr.append(spacing)

        if align:
            paragraph.alignment = align

        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            run.font.name = font_name
            run.font.bold = bold
            if font_color:
                run.font.color.rgb = font_color

def add_table_row(table):
    """Add a new row to the table with formatting."""
    tbl = table._tbl
    tr = OxmlElement('w:tr')
    tbl.append(tr)
    for _ in range(len(table.columns)):
        tc = OxmlElement('w:tc')
        tr.append(tc)
        tcPr = OxmlElement('w:tcPr')
        tc.append(tcPr)
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(table.columns[0].width))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
        p = OxmlElement('w:p')
        tc.append(p)
    return table.rows[-1]

def set_cell_borders(cell, top=False, bottom=False, left=False, right=False):
    """Set borders for a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    tcBorders = OxmlElement('w:tcBorders')
    
    border_attrs = {
        'w:val': 'single',
        'w:sz': '4',
        'w:space': '0',
        'w:color': '000000'
    }
    
    if top:
        tcBorders.append(_create_border_element('top', border_attrs))
    if bottom:
        tcBorders.append(_create_border_element('bottom', border_attrs))
    if left:
        tcBorders.append(_create_border_element('left', border_attrs))
    if right:
        tcBorders.append(_create_border_element('right', border_attrs))
    
    tcPr.append(tcBorders)

def remove_cell_borders(cell):
    """Remove all borders from a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    tcBorders = OxmlElement('w:tcBorders')
    
    border_attrs = {
        'w:val': 'nil',
        'w:sz': '0',
        'w:space': '0',
        'w:color': 'auto'
    }
    
    for side in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
        tcBorders.append(_create_border_element(side, border_attrs))
    
    tcPr.append(tcBorders)

def _create_border_element(side, attrs):
    """Helper function to create a border element."""
    border = OxmlElement(f'w:{side}')
    for key, value in attrs.items():
        border.set(qn(key), str(value))
    return border


def generate_from_excel(file_path):
    if not os.path.exists(f'output'):
        os.makedirs(f'output')
    
    try:
        df = pd.read_excel(file_path)
    except:
        df = pd.read_csv(file_path)

    log_success = []
    for inv in tqdm(df.invoice_no.unique(), total = df.invoice_no.nunique()):
        try:
            data = df[df['invoice_no'] == inv].fillna('').to_dict(orient = 'records')
            
            dt_inv = pd.to_datetime(data[0]['invoice_date']).strftime('%d/%m/%Y')

            template_path = os.path.join(os.path.dirname(__file__), 'Finance Invoice template_new.docx')
            f = open(template_path, 'rb')
            
            doc = Document(f)
            
            # Dynamic rows — line items
            subtotal = round(sum([int(x.get('amount', 0)) for x in data]))   
            vat_ori = round(sum([int(x.get('vat', 0)) for x in data]))
            total = f"Rp. {(round(subtotal + vat_ori)):,}".replace('.0', '')
            
            subtotal = f"Rp. {round(subtotal):,}".replace('.0', '')
            vat =f"Rp. {round(vat_ori):,}".replace('.0', '')
                
            inv_no = data[0].get('invoice_no')
            # Static fields

            attention = data[0].get('attention', None)

            if not attention:
                addess_company = [(14, 0, data[0].get('company_name')),
                    (15, 0, data[0].get('address_1')),
                    (16, 0, data[0].get('address_2')),
                    (17, 0, data[0].get('address_3')),
                    (18, 0, data[0].get('address_4')),
                    (19, 0, data[0].get('address_5'))]
            else:
                addess_company = [(14, 0, data[0].get('attention')),
                    (15, 0, data[0].get('company_name')),
                    (16, 0, data[0].get('address_1')),
                    (17, 0, data[0].get('address_2')),
                    (18, 0, data[0].get('address_3')),
                    (19, 0, data[0].get('address_4'))]
            
            for row, col, value in addess_company:
                set_cell_text(doc.tables[0].cell(row, col), value)

            current_row = 22

            if len(data) > 4:
                for col in [0, 2, 3, 4]:
                    remove_cell_borders(doc.tables[0].cell(25, col))
                    if col == 0:
                        set_cell_borders(doc.tables[0].cell(25, col), left = True)
                    else:
                        set_cell_borders(doc.tables[0].cell(25, col), left = True, right = True)
            for item in data:
                if current_row >= 26:
                    add_table_row(doc.tables[0])
                    for col in [0, 2, 3, 4]:
                        if col == 0:
                            set_cell_borders(doc.tables[0].cell(current_row, col), left = True)
                        else:
                            set_cell_borders(doc.tables[0].cell(current_row, col), left = True, right = True)

                set_cell_text(doc.tables[0].cell(current_row, 0), item.get('description'), align=WD_ALIGN_PARAGRAPH.LEFT)
                set_cell_text(doc.tables[0].cell(current_row, 2), f"{round(float(item.get('price_qty', ''))):,}".replace('.0', ''), align=WD_ALIGN_PARAGRAPH.RIGHT)
                set_cell_text(doc.tables[0].cell(current_row, 3), str(round(float(item.get('qty', '')))).replace('.0', ''), align=WD_ALIGN_PARAGRAPH.RIGHT)
                set_cell_text(doc.tables[0].cell(current_row, 4), f"{round(float(item.get('amount', ''))):,}".replace('.0', ''), align=WD_ALIGN_PARAGRAPH.RIGHT)
                current_row += 1
            
            if len(data) > 4:
                for col in range(5):
                    set_cell_borders(doc.tables[0].cell(current_row-1, col), bottom = True)
            else:
                current_row = 26 


            extra_rows = current_row - 26

            sub_total_row = 27 + extra_rows
            vat_row = 28 + extra_rows
            total_row = 29 + extra_rows

            bank_rows_start = 36
            bank_rows = [bank_rows_start + i + extra_rows for i in range(5)]
            
            title_note = 31 + extra_rows
            payment_notes_row = 32 + extra_rows
            payment_notes_row2 = 33 + extra_rows

            # Bold cells — subtotal, vat, total
            if vat_ori == 0:
                set_cell_text(doc.tables[0].cell(total_row, 4), str(total).replace('.0', ''),    bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
                set_cell_text(doc.tables[0].cell(total_row, 3), "Total",    bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
            else:
                set_cell_text(doc.tables[0].cell(sub_total_row, 4), str((subtotal)).replace('.0', ''), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
                set_cell_text(doc.tables[0].cell(vat_row, 4), str((vat)).replace('.0', ''),      bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
                set_cell_text(doc.tables[0].cell(total_row, 4), str((total)).replace('.0', ''),    bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

                set_cell_text(doc.tables[0].cell(sub_total_row, 3), "Subtotal", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
                set_cell_text(doc.tables[0].cell(vat_row, 3), "VAT Total",      bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
                set_cell_text(doc.tables[0].cell(total_row, 3), "Total",    bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

            static_cells = [
                (1, 4, dt_inv),
                (2, 4, inv_no),
                (3, 4, data[0].get('po')),
                (title_note, 0, 'Payment notes:'),
                (payment_notes_row, 0, data[0].get('payment_notes')),
                (payment_notes_row2, 0, '2. Please make your payment into our account bank:'),
                            ] + [
                    (bank_rows[0], 0, 'Account Name           : ' + str(data[0].get('bank_account_name', '')).replace('.0', '')),
                    (bank_rows[1], 0, 'Bank                          : ' + str(data[0].get('bank', ''))),
                    (bank_rows[2], 0, 'Account Number        : ' + str(data[0].get('bank_account_number', '')).replace('.0', '')),
                    (bank_rows[3], 0, 'Branch                       : ' + str(data[0].get('bank_branch', ''))),
                    (bank_rows[4], 0, 'Swift Code                 : ' + str(data[0].get('swift_code', ''))),
                    (bank_rows[4] + 2, 0, 'Finance Department'),
                ]

            for row, col, value in static_cells:
                set_cell_text(doc.tables[0].cell(row, col), value)

            last_row = bank_rows[4] + 2
            table = doc.tables[0]
            while len(table.rows) > last_row + 1:
                tr = table.rows[-1]._tr
                tr.getparent().remove(tr)

            tmp_name = f'output/Invoice_{inv_no}.docx'
            doc.save(tmp_name)
            f.close()
            subprocess.run(['libreoffice', '--convert-to', 'pdf' ,
                                tmp_name, '--outdir', 'output']
                            ,stdout=subprocess.DEVNULL,
                                stderr=subprocess.DEVNULL
                            )
                
            os.remove(tmp_name)
            log_success.append({'invoice_no': inv, 'status': 'Success'})
        except Exception as e:
            log_success.append({'invoice_no': inv, 'status': 'Failed', 'error': str(e).replace('\n', '  ')})
    df_log = pd.DataFrame(log_success)
    df_log.to_excel('output/log.xlsx', index=False)